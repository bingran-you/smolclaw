"""Drive bridge endpoints for mock Google Docs files."""

from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from html import escape
from io import BytesIO
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, Query, Response
from sqlalchemy.orm import Session

from claw_gdoc.models import ChangeRecord, Document, DocumentPermission, DocumentRevision, User, generate_revision_id

from .access import list_accessible_documents, require_document_access
from .deps import get_db, resolve_actor_user_id
from .history_tracker import ensure_owner_permission, record_document_change, record_revision_snapshot
from .render import default_document_style, default_named_styles, dump_json_field, render_document_resource
from .schemas import ChannelRequest, ChannelResource, DriveFileList, DriveFileResource, DriveRevisionList, DriveRevisionResource
from claw_gdoc.state.channels import channel_registry

router = APIRouter()

GOOGLE_DOCS_MIME_TYPE = "application/vnd.google-apps.document"
_QUERY_CONTAINS_RE = re.compile(r"name\s+contains\s+'([^']+)'", re.IGNORECASE)
_QUERY_FULLTEXT_RE = re.compile(r"fullText\s+contains\s+'([^']+)'", re.IGNORECASE)
_QUERY_NAME_EQUALS_RE = re.compile(r"name\s*=\s*'([^']+)'", re.IGNORECASE)
_QUERY_EQUALS_RE = re.compile(r"mimeType\s*=\s*'([^']+)'", re.IGNORECASE)
_QUERY_TRASHED_RE = re.compile(r"trashed\s*=\s*(true|false)", re.IGNORECASE)
_QUERY_ME_OWNERS_RE = re.compile(r"'me'\s+in\s+owners", re.IGNORECASE)
_QUERY_SHARED_WITH_ME_RE = re.compile(r"sharedWithMe", re.IGNORECASE)
_LIST_FIELDS_RE = re.compile(r"files\(([^)]*)\)")
_MARKDOWN_MIME_TYPES = {"text/markdown", "text/x-markdown"}


def _iso_z(value) -> str:
    return value.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _export_links(document_id: str) -> dict[str, str]:
    base = f"/drive/v3/files/{document_id}/export?mimeType="
    return {
        "application/pdf": f"{base}application/pdf",
        "application/rtf": f"{base}application/rtf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": (
            f"{base}application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        "text/html": f"{base}text/html",
        "text/markdown": f"{base}text/markdown",
        "text/plain": f"{base}text/plain",
        "text/x-markdown": f"{base}text/x-markdown",
    }


def _user_resource(user: User | None) -> dict[str, object] | None:
    if user is None:
        return None
    return {
        "displayName": user.display_name,
        "emailAddress": user.email_address,
        "kind": "drive#user",
        "me": False,
        "permissionId": user.id,
    }


def _drive_file_payload(document: Document, *, owned_by_me: bool = True) -> dict[str, object]:
    owner_resource = _user_resource(document.user)
    shared = any(
        permission.role != "owner" or permission.user_id != document.user_id
        for permission in document.permissions
    )
    return {
        "kind": "drive#file",
        "id": document.id,
        "name": document.title,
        "mimeType": GOOGLE_DOCS_MIME_TYPE,
        "description": document.description or None,
        "createdTime": _iso_z(document.created_at),
        "modifiedTime": _iso_z(document.updated_at),
        "trashed": document.trashed,
        "ownedByMe": owned_by_me,
        "webViewLink": f"https://docs.google.com/document/d/{document.id}/edit?usp=drivesdk",
        "iconLink": "https://drive-thirdparty.googleusercontent.com/16/type/application/vnd.google-apps.document",
        "exportLinks": _export_links(document.id),
        "shared": shared,
        "version": str(max(1, len(document.revisions))),
        "headRevisionId": str(document.revision_id),
        "size": str(len(document.body_text.encode("utf-8"))),
        "owners": [owner_resource] if owner_resource else None,
        "lastModifyingUser": owner_resource,
    }


def _requested_get_fields(fields: str | None) -> set[str]:
    if not fields:
        return {"kind", "id", "name", "mimeType", "ownedByMe", "headRevisionId", "version", "shared"}
    return {field.strip() for field in fields.split(",") if field.strip()}


def _requested_list_fields(fields: str | None) -> tuple[set[str], set[str]]:
    if not fields:
        return {"kind", "files", "nextPageToken"}, {
            "kind",
            "id",
            "name",
            "mimeType",
            "ownedByMe",
            "headRevisionId",
            "version",
            "shared",
        }
    top_fields = {
        field.strip()
        for field in _LIST_FIELDS_RE.sub("", fields).split(",")
        if field.strip()
    }
    match = _LIST_FIELDS_RE.search(fields)
    file_fields = {"kind", "id", "name", "mimeType", "ownedByMe"}
    if match:
        top_fields.add("files")
        file_fields = {field.strip() for field in match.group(1).split(",") if field.strip()}
    return top_fields or {"kind", "files"}, file_fields


def _drive_file_resource(
    document: Document,
    *,
    fields: str | None = None,
    owned_by_me: bool = True,
) -> DriveFileResource:
    payload = _drive_file_payload(document, owned_by_me=owned_by_me)
    selected = {field: payload[field] for field in _requested_get_fields(fields) if field in payload}
    return DriveFileResource.model_validate(selected)


def _revision_resource(revision: DocumentRevision, *, user: User | None = None) -> DriveRevisionResource:
    export_links = json.loads(revision.export_links_json) if revision.export_links_json else _export_links(revision.document_id)
    return DriveRevisionResource(
        id=revision.revision_id,
        modifiedTime=_iso_z(revision.created_at),
        keepForever=revision.keep_forever,
        size=str(revision.file_size),
        exportLinks=export_links,
        lastModifyingUser=_user_resource(user),
        originalFilename=f"{revision.title or 'Untitled document'}.gdoc",
    )


def _resolve_document(
    db: Session,
    actor_user_id: str,
    file_id: str,
    *,
    minimum_role: str = "reader",
) -> tuple[Document, bool]:
    document = db.query(Document).filter(Document.id == file_id).first()
    document, permission = require_document_access(
        db,
        document,
        actor_user_id,
        minimum_role=minimum_role,
        not_found_message="File not found",
    )
    owned_by_me = permission.role == "owner" and permission.user_id == document.user_id
    return document, owned_by_me


def _match_query(document: Document, query: str, *, owned_by_me: bool) -> bool:
    query = query.strip()
    if not query:
        return True
    lowered = query.lower()
    if "and" in lowered:
        return all(
            _match_query(document, part.strip(), owned_by_me=owned_by_me)
            for part in re.split(r"\band\b", query, flags=re.IGNORECASE)
        )

    contains_match = _QUERY_CONTAINS_RE.search(query)
    if contains_match:
        return contains_match.group(1).lower() in document.title.lower()

    fulltext_match = _QUERY_FULLTEXT_RE.search(query)
    if fulltext_match:
        return fulltext_match.group(1).lower() in document.body_text.lower()

    name_equals_match = _QUERY_NAME_EQUALS_RE.fullmatch(query)
    if name_equals_match:
        return document.title == name_equals_match.group(1)

    mime_match = _QUERY_EQUALS_RE.fullmatch(query)
    if mime_match:
        return mime_match.group(1) == GOOGLE_DOCS_MIME_TYPE

    trashed_match = _QUERY_TRASHED_RE.fullmatch(query)
    if trashed_match:
        return document.trashed is (trashed_match.group(1).lower() == "true")

    if _QUERY_ME_OWNERS_RE.fullmatch(query):
        return owned_by_me

    if _QUERY_SHARED_WITH_ME_RE.fullmatch(query):
        return not owned_by_me

    return query.lower() in document.title.lower()


def _sort_documents(
    documents_with_permissions: list[tuple[Document, object]],
    order_by: str | None,
) -> list[tuple[Document, object]]:
    if not order_by:
        return sorted(
            documents_with_permissions,
            key=lambda item: (item[0].updated_at, item[0].id),
            reverse=True,
        )

    order_terms = [term.strip() for term in order_by.split(",") if term.strip()]
    ordered = list(documents_with_permissions)
    for term in reversed(order_terms):
        descending = term.endswith(" desc")
        key_name = term[:-5] if descending else term
        key_name = key_name.replace(" asc", "").strip()
        if key_name == "name":
            key_fn = lambda item: item[0].title.lower()
        elif key_name in {"modifiedTime", "modifiedDate"}:
            key_fn = lambda item: item[0].updated_at
        elif key_name == "createdTime":
            key_fn = lambda item: item[0].created_at
        elif key_name == "name_natural":
            key_fn = lambda item: item[0].title
        else:
            continue
        ordered.sort(key=key_fn, reverse=descending)
    return ordered


def _document_text_to_html(text: str) -> str:
    lines = text.rstrip("\n").split("\n")
    blocks = []
    for line in lines:
        if line.startswith("- "):
            blocks.append(f"<li>{line[2:]}</li>")
        else:
            blocks.append(f"<p>{line}</p>")
    html = []
    in_list = False
    for block in blocks:
        if block.startswith("<li>"):
            if not in_list:
                html.append("<ul>")
                in_list = True
            html.append(block)
        else:
            if in_list:
                html.append("</ul>")
                in_list = False
            html.append(block)
    if in_list:
        html.append("</ul>")
    return "".join(html)


def _document_text_to_markdown(text: str) -> str:
    return text


def _document_resource_for_export(document: Document):
    return render_document_resource(
        document_id=document.id,
        title=document.title,
        body_text=document.body_text,
        text_style_spans_json=document.text_style_spans_json,
        paragraph_style_json=document.paragraph_style_json,
        named_ranges_json=document.named_ranges_json,
        named_styles_json=document.named_styles_json,
        document_style_json=document.document_style_json,
        revision_id=document.revision_id,
    )


def _run_to_html(text: str, text_style: dict) -> str:
    rendered = escape(text)
    link = (text_style.get("link") or {}).get("url")
    if link:
        rendered = f'<a href="{escape(link, quote=True)}">{rendered}</a>'
    if text_style.get("bold"):
        rendered = f"<strong>{rendered}</strong>"
    if text_style.get("italic"):
        rendered = f"<em>{rendered}</em>"
    if text_style.get("underline"):
        rendered = f"<u>{rendered}</u>"
    return rendered


def _run_to_markdown(text: str, text_style: dict) -> str:
    rendered = text
    link = (text_style.get("link") or {}).get("url")
    if link:
        rendered = f"[{rendered}]({link})"
    if text_style.get("underline"):
        rendered = f"<u>{rendered}</u>"
    if text_style.get("italic"):
        rendered = f"*{rendered}*"
    if text_style.get("bold"):
        rendered = f"**{rendered}**"
    return rendered


def _document_to_html(document: Document) -> str:
    resource = _document_resource_for_export(document)
    parts: list[str] = []
    in_list = False
    for element in resource.body.content[1:]:
        paragraph = element.paragraph
        if paragraph is None:
            continue
        runs = "".join(
            _run_to_html(
                run.textRun.content.rstrip("\n"),
                run.textRun.textStyle.model_dump(exclude_none=True),
            )
            for run in paragraph.elements
        )
        named_style = paragraph.paragraphStyle.namedStyleType
        is_bullet = paragraph.bullet is not None
        if is_bullet and not in_list:
            parts.append("<ul>")
            in_list = True
        if not is_bullet and in_list:
            parts.append("</ul>")
            in_list = False
        if named_style == "HEADING_1":
            parts.append(f"<h1>{runs}</h1>")
        elif named_style == "HEADING_2":
            parts.append(f"<h2>{runs}</h2>")
        elif named_style == "HEADING_3":
            parts.append(f"<h3>{runs}</h3>")
        elif is_bullet:
            parts.append(f"<li>{runs}</li>")
        else:
            parts.append(f"<p>{runs}</p>")
    if in_list:
        parts.append("</ul>")
    return "".join(parts)


def _document_to_markdown(document: Document) -> str:
    resource = _document_resource_for_export(document)
    lines: list[str] = []
    for element in resource.body.content[1:]:
        paragraph = element.paragraph
        if paragraph is None:
            continue
        text = "".join(
            _run_to_markdown(
                run.textRun.content.rstrip("\n"),
                run.textRun.textStyle.model_dump(exclude_none=True),
            )
            for run in paragraph.elements
        )
        named_style = paragraph.paragraphStyle.namedStyleType
        if named_style == "HEADING_1":
            lines.append(f"# {text}")
        elif named_style == "HEADING_2":
            lines.append(f"## {text}")
        elif named_style == "HEADING_3":
            lines.append(f"### {text}")
        elif paragraph.bullet is not None:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    return "\n".join(lines).rstrip() + "\n"


def _escape_rtf(text: str) -> str:
    return text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")


def _document_text_to_rtf(text: str) -> str:
    parts = ["{\\rtf1\\ansi\n"]
    for line in text.rstrip("\n").split("\n"):
        parts.append(_escape_rtf(line))
        parts.append("\\par\n")
    parts.append("}")
    return "".join(parts)


def _document_text_to_docx(text: str) -> bytes:
    document = DocxDocument()
    for line in text.rstrip("\n").split("\n"):
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _document_to_docx(document: Document) -> bytes:
    resource = _document_resource_for_export(document)
    docx_document = DocxDocument()
    for element in resource.body.content[1:]:
        paragraph = element.paragraph
        if paragraph is None:
            continue
        style_name = None
        if paragraph.bullet is not None:
            style_name = "List Bullet"
        elif paragraph.paragraphStyle.namedStyleType == "HEADING_1":
            style_name = "Heading 1"
        elif paragraph.paragraphStyle.namedStyleType == "HEADING_2":
            style_name = "Heading 2"
        elif paragraph.paragraphStyle.namedStyleType == "HEADING_3":
            style_name = "Heading 3"
        docx_paragraph = docx_document.add_paragraph(style=style_name)
        for element_run in paragraph.elements:
            content = element_run.textRun.content.rstrip("\n")
            run = docx_paragraph.add_run(content)
            style = element_run.textRun.textStyle
            if style.bold:
                run.bold = True
            if style.italic:
                run.italic = True
            if style.underline:
                run.underline = True
    buffer = BytesIO()
    docx_document.save(buffer)
    return buffer.getvalue()


def _document_text_to_pdf_bytes(text: str) -> bytes:
    lines = text.rstrip("\n").split("\n") or [""]
    escaped_lines = [
        line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        for line in lines
    ]
    content_lines = ["BT", "/F1 12 Tf", "72 720 Td", "14 TL"]
    first = True
    for line in escaped_lines:
        if first:
            content_lines.append(f"({line}) Tj")
            first = False
        else:
            content_lines.append(f"T* ({line}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("utf-8")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(stream)} >>\nstream\n".encode("utf-8") + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{idx} 0 obj\n".encode("utf-8"))
        result.extend(obj)
        result.extend(b"\nendobj\n")

    xref_start = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode("utf-8"))
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode("utf-8"))
    result.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF"
        ).encode("utf-8")
    )
    return bytes(result)


def _new_document(
    *,
    actor_user_id: str,
    name: str,
    description: str,
    body_text: str = "\n",
) -> Document:
    now = datetime.now(timezone.utc)
    return Document(
        id=uuid4().hex[:32],
        user_id=actor_user_id,
        title=name or "Untitled document",
        description=description,
        body_text=body_text,
        text_style_spans_json="[]",
        paragraph_style_json="[]",
        named_ranges_json="[]",
        named_styles_json=dump_json_field(default_named_styles()),
        document_style_json=dump_json_field(default_document_style()),
        revision_id=generate_revision_id(),
        created_at=now,
        updated_at=now,
        trashed=False,
    )


@router.get("/files", response_model=DriveFileList, response_model_exclude_none=True)
def list_files(
    q: str | None = Query(None),
    fields: str | None = Query(None),
    orderBy: str | None = Query(None),
    pageSize: int = Query(100, ge=1, le=1000),
    pageToken: str | None = Query(None),
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    documents_with_permissions = list_accessible_documents(db, actor_user_id)
    if q:
        documents_with_permissions = [
            (document, permission)
            for document, permission in documents_with_permissions
            if _match_query(
                document,
                q,
                owned_by_me=(permission.role == "owner" and permission.user_id == document.user_id),
            )
        ]
    documents_with_permissions = _sort_documents(documents_with_permissions, orderBy)

    offset = 0
    if pageToken:
        try:
            offset = int(pageToken)
        except ValueError as exc:
            raise HTTPException(
                400,
                {"message": "Invalid pageToken", "reason": "badRequest"},
            ) from exc

    sliced = documents_with_permissions[offset : offset + pageSize]
    next_page_token = str(offset + pageSize) if offset + pageSize < len(documents_with_permissions) else None
    top_fields, file_fields = _requested_list_fields(fields)

    payload: dict[str, object] = {}
    if "kind" in top_fields:
        payload["kind"] = "drive#fileList"
    if "files" in top_fields:
        payload["files"] = [
            DriveFileResource.model_validate(
                {
                    field: value
                    for field, value in _drive_file_payload(
                        document,
                        owned_by_me=(permission.role == "owner" and permission.user_id == document.user_id),
                    ).items()
                    if field in file_fields
                }
            )
            for document, permission in sliced
        ]
    if "nextPageToken" in top_fields and next_page_token is not None:
        payload["nextPageToken"] = next_page_token
    return DriveFileList.model_validate(payload)


@router.get("/files/{fileId}", response_model=DriveFileResource, response_model_exclude_none=True)
def get_file(
    fileId: str,
    fields: str | None = Query(None),
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    document, owned_by_me = _resolve_document(db, actor_user_id, fileId)
    return _drive_file_resource(document, fields=fields, owned_by_me=owned_by_me)


@router.post("/files", response_model=DriveFileResource, response_model_exclude_none=True)
def create_file(
    body: dict,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    mime_type = body.get("mimeType") or GOOGLE_DOCS_MIME_TYPE
    if mime_type != GOOGLE_DOCS_MIME_TYPE:
        raise HTTPException(
            400,
            {"message": f"Unsupported mimeType: {mime_type}", "reason": "badRequest"},
        )
    document = _new_document(
        actor_user_id=actor_user_id,
        name=body.get("name") or body.get("title") or "Untitled document",
        description=body.get("description", ""),
    )
    db.add(document)
    db.flush()
    ensure_owner_permission(db, document)
    record_revision_snapshot(db, document, actor_user_id=actor_user_id)
    record_document_change(db, document, change_type="fileCreated")
    db.commit()
    db.refresh(document)
    return _drive_file_resource(document)


@router.post("/files/{fileId}/copy", response_model=DriveFileResource, response_model_exclude_none=True)
def copy_file(
    fileId: str,
    body: dict,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    source, _owned_by_me = _resolve_document(db, actor_user_id, fileId)
    now = datetime.now(timezone.utc)
    clone = Document(
        id=uuid4().hex[:32],
        user_id=actor_user_id,
        title=body.get("name") or source.title,
        description=body.get("description", source.description),
        body_text=source.body_text,
        text_style_spans_json=source.text_style_spans_json,
        paragraph_style_json=source.paragraph_style_json,
        named_ranges_json=source.named_ranges_json,
        named_styles_json=source.named_styles_json,
        document_style_json=source.document_style_json,
        revision_id=generate_revision_id(),
        created_at=now,
        updated_at=now,
        trashed=False,
    )
    db.add(clone)
    db.flush()
    ensure_owner_permission(db, clone)
    record_revision_snapshot(db, clone, actor_user_id=actor_user_id)
    record_document_change(db, clone, change_type="fileCreated")
    db.commit()
    db.refresh(clone)
    return _drive_file_resource(clone)


def _update_file_common(
    *,
    fileId: str,
    body: dict,
    db: Session,
    actor_user_id: str,
):
    document, owned_by_me = _resolve_document(db, actor_user_id, fileId, minimum_role="writer")
    if "name" in body:
        document.title = body["name"] or document.title
    if "description" in body:
        document.description = body.get("description") or ""
    if "trashed" in body:
        document.trashed = bool(body["trashed"])
    document.updated_at = datetime.now(timezone.utc)
    document.revision_id = generate_revision_id()
    record_revision_snapshot(db, document, actor_user_id=actor_user_id)
    record_document_change(db, document, change_type="fileUpdated")
    db.commit()
    db.refresh(document)
    return _drive_file_resource(document, owned_by_me=owned_by_me)


@router.patch("/files/{fileId}", response_model=DriveFileResource, response_model_exclude_none=True)
def update_file(
    fileId: str,
    body: dict,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    return _update_file_common(
        fileId=fileId,
        body=body,
        db=db,
        actor_user_id=actor_user_id,
    )


@router.put("/files/{fileId}", response_model=DriveFileResource, response_model_exclude_none=True)
def replace_file(
    fileId: str,
    body: dict,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    return _update_file_common(
        fileId=fileId,
        body=body,
        db=db,
        actor_user_id=actor_user_id,
    )


@router.delete("/files/{fileId}", status_code=204)
def delete_file(
    fileId: str,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    document, _owned_by_me = _resolve_document(db, actor_user_id, fileId, minimum_role="owner")
    record_document_change(
        db,
        document,
        change_type="fileDeleted",
        removed=True,
        removed_file_payload=_drive_file_payload(document, owned_by_me=True),
    )
    (
        db.query(ChangeRecord)
        .filter(ChangeRecord.file_id == document.id)
        .update({"document_id": None}, synchronize_session=False)
    )
    db.delete(document)
    db.commit()
    return Response(status_code=204)


@router.get("/files/{fileId}/revisions", response_model=DriveRevisionList, response_model_exclude_none=True)
def list_revisions(
    fileId: str,
    pageSize: int = Query(100, ge=1, le=1000),
    pageToken: str | None = Query(None),
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    document, _owned_by_me = _resolve_document(db, actor_user_id, fileId)
    query = (
        db.query(DocumentRevision)
        .filter(DocumentRevision.document_id == document.id)
        .order_by(DocumentRevision.created_at.desc(), DocumentRevision.id.desc())
    )
    offset = 0
    if pageToken:
        try:
            offset = int(pageToken)
        except ValueError as exc:
            raise HTTPException(400, {"message": "Invalid pageToken", "reason": "badRequest"}) from exc

    revisions = query.offset(offset).limit(pageSize + 1).all()
    has_more = len(revisions) > pageSize
    items = revisions[:pageSize]
    return DriveRevisionList(
        revisions=[
            _revision_resource(
                revision,
                user=db.query(User).filter(User.id == revision.user_id).first(),
            )
            for revision in items
        ],
        nextPageToken=str(offset + pageSize) if has_more else None,
    )


@router.get("/files/{fileId}/revisions/{revisionId}", response_model=DriveRevisionResource, response_model_exclude_none=True)
def get_revision(
    fileId: str,
    revisionId: str,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    document, _owned_by_me = _resolve_document(db, actor_user_id, fileId)
    revision = (
        db.query(DocumentRevision)
        .filter(
            DocumentRevision.document_id == document.id,
            DocumentRevision.revision_id == revisionId,
        )
        .first()
    )
    if revision is None:
        raise HTTPException(404, "Revision not found")
    user = db.query(User).filter(User.id == revision.user_id).first()
    return _revision_resource(revision, user=user)


@router.get("/files/{fileId}/export")
def export_file(
    fileId: str,
    mimeType: str = Query(...),
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    document, _owned_by_me = _resolve_document(db, actor_user_id, fileId)
    if mimeType == "text/plain":
        return Response(content=document.body_text, media_type="text/plain")
    if mimeType == "text/html":
        return Response(content=_document_to_html(document), media_type="text/html")
    if mimeType in _MARKDOWN_MIME_TYPES:
        return Response(content=_document_to_markdown(document), media_type=mimeType)
    if mimeType == "application/rtf":
        return Response(content=_document_text_to_rtf(document.body_text), media_type=mimeType)
    if mimeType == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return Response(content=_document_to_docx(document), media_type=mimeType)
    if mimeType == "application/pdf":
        return Response(content=_document_text_to_pdf_bytes(document.body_text), media_type=mimeType)
    raise HTTPException(
        400,
        {"message": f"Unsupported export mimeType: {mimeType}", "reason": "badRequest"},
    )


@router.post("/files/{fileId}/watch", response_model=ChannelResource, response_model_exclude_none=True)
def watch_file(
    fileId: str,
    body: ChannelRequest,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    _document, _owned_by_me = _resolve_document(db, actor_user_id, fileId)
    resource_uri = f"/drive/v3/files/{fileId}"
    channel = channel_registry.register(
        resource_uri=resource_uri,
        channel_id=body.id,
        address=body.address,
        token=body.token,
        channel_type=body.type,
        payload=body.payload,
        params=body.params,
    )
    return ChannelResource.model_validate(channel)


@router.post("/changes/watch", response_model=ChannelResource, response_model_exclude_none=True)
def watch_changes(
    body: ChannelRequest,
    db: Session = Depends(get_db),
    actor_user_id: str = Depends(resolve_actor_user_id),
):
    resource_uri = f"/drive/v3/changes?user={actor_user_id}"
    channel = channel_registry.register(
        resource_uri=resource_uri,
        channel_id=body.id,
        address=body.address,
        token=body.token,
        channel_type=body.type,
        payload=body.payload,
        params=body.params,
    )
    return ChannelResource.model_validate(channel)


@router.post("/channels/stop", status_code=204)
def stop_channel(body: ChannelRequest):
    if not body.id or not body.resourceId:
        raise HTTPException(400, {"message": "Missing required field: id/resourceId", "reason": "badRequest"})
    stopped = channel_registry.stop(body.id, body.resourceId)
    if not stopped:
        raise HTTPException(404, "Channel not found")
    return Response(status_code=204)
